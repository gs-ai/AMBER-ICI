"""
File Ingestion System
Handles PDFs, images, screenshots, text files, and transcripts
"""

import os
from pathlib import Path
from typing import Dict, Optional
import asyncio

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    from PIL import Image
    import pytesseract
except ImportError:
    Image = None
    pytesseract = None

try:
    import docx
except ImportError:
    docx = None


class FileProcessor:
    """Process various file types and extract text content"""
    
    def __init__(self):
        self.supported_types = {
            '.pdf': self._process_pdf,
            '.txt': self._process_text,
            '.md': self._process_text,
            '.png': self._process_image,
            '.jpg': self._process_image,
            '.jpeg': self._process_image,
            '.gif': self._process_image,
            '.bmp': self._process_image,
            '.docx': self._process_docx,
        }
    
    async def process_file(self, file_path: Path) -> Dict:
        """Process a file and extract content"""
        if not file_path.exists():
            return {"error": "File not found", "success": False}
        
        file_ext = file_path.suffix.lower()
        
        if file_ext not in self.supported_types:
            return {
                "error": f"Unsupported file type: {file_ext}",
                "success": False,
                "supported_types": list(self.supported_types.keys())
            }
        
        try:
            processor = self.supported_types[file_ext]
            result = await processor(file_path)
            result["filename"] = file_path.name
            result["file_type"] = file_ext
            result["success"] = True
            return result
        
        except Exception as e:
            return {
                "error": str(e),
                "filename": file_path.name,
                "file_type": file_ext,
                "success": False
            }
    
    async def _process_pdf(self, file_path: Path) -> Dict:
        """Extract text from PDF"""
        if PyPDF2 is None:
            return {"error": "PyPDF2 not installed", "text": ""}
        
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                num_pages = len(pdf_reader.pages)
                
                for page_num in range(num_pages):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n\n"
            
            return {
                "text": text.strip(),
                "pages": num_pages,
                "method": "pdf_extraction"
            }
        
        except Exception as e:
            return {"error": str(e), "text": ""}
    
    async def _process_text(self, file_path: Path) -> Dict:
        """Read plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            return {
                "text": text,
                "lines": len(text.split('\n')),
                "method": "text_read"
            }
        
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as file:
                text = file.read()
            
            return {
                "text": text,
                "lines": len(text.split('\n')),
                "method": "text_read",
                "encoding": "latin-1"
            }
    
    async def _process_image(self, file_path: Path) -> Dict:
        """Extract text from image using OCR"""
        if Image is None or pytesseract is None:
            return {
                "error": "PIL or pytesseract not installed",
                "text": "",
                "method": "ocr_unavailable"
            }
        
        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)
            
            return {
                "text": text,
                "image_size": image.size,
                "image_mode": image.mode,
                "method": "ocr"
            }
        
        except Exception as e:
            return {
                "error": str(e),
                "text": "",
                "method": "ocr_failed"
            }
    
    async def _process_docx(self, file_path: Path) -> Dict:
        """Extract text from Word document"""
        if docx is None:
            return {"error": "python-docx not installed", "text": ""}
        
        try:
            doc = docx.Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            return {
                "text": text,
                "paragraphs": len(doc.paragraphs),
                "method": "docx_extraction"
            }
        
        except Exception as e:
            return {"error": str(e), "text": ""}
